"""
Document Upload and Basic Parsing Module
Milestone 1 - Week 1-2
"""

import PyPDF2
from docx import Document
from pathlib import Path
from typing import Dict, List


class DocumentParser:
    """Parse legal contracts from PDF and Word documents"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.file_type = Path(file_path).suffix.lower()
        self.text = ""
        self.metadata = {}
        
    def parse(self) -> Dict:
        """
        Parse document and extract text content
        Returns: Dictionary with text and metadata
        """
        if self.file_type == '.pdf':
            return self._parse_pdf()
        elif self.file_type in ['.docx', '.doc']:
            return self._parse_word()
        else:
            raise ValueError(f"Unsupported file type: {self.file_type}")
    
    def _parse_pdf(self) -> Dict:
        """Parse PDF document"""
        try:
            with open(self.file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                self.metadata = {
                    'page_count': len(pdf_reader.pages),
                    'file_type': 'PDF',
                    'file_name': Path(self.file_path).name
                }
                
                # Extract text from all pages
                text_parts = []
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {page_num} ---\n{page_text}")
                
                self.text = "\n\n".join(text_parts)
                
                return {
                    'text': self.text,
                    'metadata': self.metadata
                }
                
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
    
    def _parse_word(self) -> Dict:
        """Parse Word document"""
        try:
            doc = Document(self.file_path)
            
            # Extract metadata
            self.metadata = {
                'paragraph_count': len(doc.paragraphs),
                'file_type': 'Word',
                'file_name': Path(self.file_path).name
            }
            
            # Extract text from all paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            self.text = "\n\n".join(text_parts)
            
            return {
                'text': self.text,
                'metadata': self.metadata
            }
            
        except Exception as e:
            raise Exception(f"Error parsing Word document: {str(e)}")
    
    def get_text_chunks(self, chunk_size: int = 1000) -> List[str]:
        """
        Split text into chunks for vector embedding
        Used for Pinecone vector database
        """
        words = self.text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size):
            chunk = ' '.join(words[i:i + chunk_size])
            chunks.append(chunk)
        
        return chunks


if __name__ == "__main__":
    # Test the parser
    print("Document Parser - Milestone 1")
    print("=" * 60)
    
    # Test with sample Word document
    test_file = "sample_contract.docx"
    if Path(test_file).exists():
        parser = DocumentParser(test_file)
        result = parser.parse()
        
        print(f"\nFile: {result['metadata']['file_name']}")
        print(f"Type: {result['metadata']['file_type']}")
        print(f"Text length: {len(result['text'])} characters")
        print(f"\nFirst 200 characters:")
        print(result['text'][:200])
        
        # Test chunking for vector database
        chunks = parser.get_text_chunks(chunk_size=500)
        print(f"\nText chunks for Pinecone: {len(chunks)} chunks")
    else:
        print(f"Test file '{test_file}' not found")
