from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import markdown

class ClauseAIPDF(FPDF):
    def header(self):
        # Only print header on pages > 1 to avoid title duplication on cover
        if self.page_no() > 1:
            self.set_font('helvetica', 'I', 8)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, 'ClauseAI Analysis Report', align='R')
            self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font('helvetica', 'I', 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f'Page {self.page_no()}', align='C')

    def chapter_body(self, text):
        self.set_font('helvetica', size=11)
        self.set_text_color(0, 0, 0)
        
        # Clean text
        replacements = {
            "’": "'", "‘": "'", "“": '"', "”": '"', "–": "-", "—": "-",
            "…": "...", "₹": "Rs.", "€": "EUR", "£": "GBP",
            "✅": "[PASS]", "⚠️": "[RISK]", "❌": "[FAIL]",
            "●": "-", "•": "-"
        }
        for char, replacement in replacements.items():
            text = text.replace(char, replacement)
        
        # Encode
        text = text.encode('latin-1', 'replace').decode('latin-1')
        
        lines = text.split('\n')
        
        # Title Handling: If the first line is a title, make it big
        first_line = True
        
        for line in lines:
            line = line.strip()
            if not line:
                self.ln(2)
                continue
                
            # Handle Headers
            if line.startswith('#'):
                level = len(line.split(' ')[0])
                clean_line = line.lstrip('#').strip()
                
                # Special case: Title detection (first line or #)
                if level == 1:
                    self.ln(5)
                    self.set_font('helvetica', 'B', 16)
                    self.set_fill_color(240, 240, 240)
                    self.cell(0, 10, clean_line, ln=True, fill=True)
                    self.set_font('helvetica', size=11)
                elif level == 2:
                    self.ln(4)
                    self.set_font('helvetica', 'B', 14)
                    self.cell(0, 8, clean_line, ln=True)
                    self.set_font('helvetica', size=11)
                elif level == 3:
                    self.ln(2)
                    self.set_font('helvetica', 'B', 12)
                    self.cell(0, 6, clean_line, ln=True)
                    self.set_font('helvetica', size=11)
                else:
                    self.set_font('helvetica', 'B', 11)
                    self.cell(0, 6, clean_line, ln=True)
                    self.set_font('helvetica', size=11)

            # Handle Bullet Points
            elif line.startswith('- ') or line.startswith('* '):
                self.set_x(self.l_margin + 5)
                bullet_text = line[2:]
                self._write_markdown_line(bullet_text)
                
            # Standard Text
            else:
                self._write_markdown_line(line)
                
    def _write_markdown_line(self, text):
        """Helpers to handle bold text (**text**) within a line"""
        if "**" not in text:
            self.multi_cell(0, 5, text)
            self.ln(1)
            return

        parts = text.split("**")
        for i, part in enumerate(parts):
            if i % 2 == 1: # Odd = Bold
                self.set_font('helvetica', 'B', 11)
                self.write(5, part)
            else:
                self.set_font('helvetica', '', 11)
                self.write(5, part)
        self.ln(6) # End of line

def generate_pdf(report_text: str) -> bytes:
    """
    Generates a professionally formatted PDF report.
    """
    pdf = ClauseAIPDF()
    pdf.add_page()
    
    # Generate the main body text of the report using FPDF
    pdf.chapter_body(report_text)
    return bytes(pdf.output())

def add_markdown_paragraph(doc, text, style=None):
    """
    Helper to add a paragraph with bold (**text**) support to a python-docx Document.
    """
    p = doc.add_paragraph(style=style)
    
    if "**" not in text:
        p.add_run(text)
        return

    parts = text.split("**")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        if i % 2 == 1: # Odd indices are between ** **, so make them bold
            run.bold = True

def generate_word(report_text: str) -> bytes:
    """
    Generates a Word (DOCX) file from the given report text with beautiful formatting.
    """
    doc = Document()
    
    # Core aesthetic styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title formatting
    title = doc.add_heading('ClauseAI Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_font = title.runs[0].font
    title_font.color.rgb = RGBColor(16, 185, 129) # Emerald Green
    title_font.name = 'Segoe UI'
    
    doc.add_paragraph() # Spacer
    
    # Set default heading styles for a cleaner look
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Segoe UI'
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor(15, 23, 42) # Slate 800
    
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Segoe UI'
    h2.font.size = Pt(14)
    h2.font.color.rgb = RGBColor(51, 65, 85) # Slate 700
    
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Segoe UI'
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
    
    # Sanitize characters to prevent docx encoding errors
    replacements = {
        "’": "'", "‘": "'", "“": '"', "”": '"', "–": "-", "—": "-",
        "…": "...", "₹": "Rs.", "€": "EUR", "£": "GBP",
        "✅": "[PASS]", "⚠️": "[RISK]", "❌": "[FAIL]",
        "●": "-", "•": "-"
    }
    for char, replacement in replacements.items():
        report_text = report_text.replace(char, replacement)

    # Manually parse markdown tags since docx doesn't support them natively
    for line in report_text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph() # Add visual padding for empty lines
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point with bold support
            add_markdown_paragraph(doc, line[2:], style='List Bullet')
        else:
            # Standard paragraph with bold support
            add_markdown_paragraph(doc, line)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generate_html(report_text: str) -> bytes:
    """
    Generates a beautifully styled Web Report (HTML) as a robust alternative to PDF for non-Latin scripts.
    Automatically handles any Unicode structure seamlessly.
    """
    # Parse the markdown safely
    md_html = markdown.markdown(report_text)
    
    html_content = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>ClauseAI Analysis Report</title>
        <style>
            :root {{
                --primary: #10b981;
                --bg: #f8fafc;
                --text: #1e293b;
                --card: #ffffff;
            }}
            body {{
                font-family: 'Segoe UI', system-ui, -apple-system, sans-serif;
                background-color: var(--bg);
                color: var(--text);
                line-height: 1.6;
                padding: 40px 20px;
                max-width: 900px;
                margin: 0 auto;
            }}
            .report-container {{
                background-color: var(--card);
                padding: 40px 50px;
                border-radius: 12px;
                box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
                border-top: 5px solid var(--primary);
            }}
            h1, h2, h3, h4 {{
                color: #0f172a;
                margin-top: 1.5em;
                margin-bottom: 0.5em;
            }}
            h1 {{
                text-align: center;
                color: var(--primary);
                font-size: 2.5em;
                border-bottom: 2px solid #e2e8f0;
                padding-bottom: 20px;
                margin-bottom: 30px;
                margin-top: 0;
            }}
            h2 {{
                font-size: 1.75em;
                border-bottom: 1px solid #e2e8f0;
                padding-bottom: 10px;
            }}
            p {{ margin-bottom: 1em; }}
            ul, ol {{ margin-bottom: 1.5em; padding-left: 2em; }}
            li {{ margin-bottom: 0.5em; }}
            strong {{ color: #0f172a; font-weight: 600; }}
            @media print {{
                body {{ background-color: white; padding: 0; }}
                .report-container {{ box-shadow: none; border: none; padding: 0; }}
            }}
        </style>
    </head>
    <body>
        <div class="report-container">
            <h1>ClauseAI Analysis Report</h1>
            {md_html}
            <div style="margin-top: 50px; text-align: center; font-size: 0.8em; color: #64748b;">
                Generated by ClauseAI Intelligent Contract Analysis System
            </div>
        </div>
    </body>
    </html>
    """
    return html_content.encode('utf-8')
