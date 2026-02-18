from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import ast
import re

# --- 1. DEEP CLEANING HELPER (The Missing Piece) ---
def clean_for_docx(raw_data):
    """
    Cleans AI output for Word Docs. 
    Removes JSON wrappers, signatures, and Markdown formatting.
    """
    text = str(raw_data).strip()
    
    # A. Parse Dictionary/JSON Structure
    try:
        # If it looks like a list or dict, parse it
        if (text.startswith("{") and text.endswith("}")) or (text.startswith("[") and text.endswith("]")):
            parsed = ast.literal_eval(text)
            if isinstance(parsed, list) and len(parsed) > 0:
                text = parsed[0].get('text', str(parsed[0])) if isinstance(parsed[0], dict) else str(parsed[0])
            elif isinstance(parsed, dict):
                text = parsed.get('text', str(parsed))
    except:
        pass # Fallback to regex if parsing fails

    # B. Brute Force Cleanup (for stubborn artifacts)
    if "{'type': 'text'" in text:
        start = text.find("'text': '")
        if start != -1: text = text[start + 9:]
    
    # Remove signatures/extras
    extras_markers = ["', 'extras':", '", "extras":', "', 'type':"]
    for marker in extras_markers:
        if marker in text: text = text.split(marker)[0]

    # C. Formatting Cleanup
    text = text.replace("\\n", "\n").replace("\\'", "'").replace('"', '"').replace("'", "'")
    text = text.rstrip("'\"}]")
    
    # D. Remove Markdown (Word doesn't support **bold** syntax directly)
    # We strip the asterisks so it looks like clean plain text
    text = text.replace("**", "").replace("###", "").replace("##", "")
    
    return text.strip()

# --- 2. DOCX GENERATOR ---
def generate_docx(results, config):
    doc = Document()
    
    # Title
    title = doc.add_heading('CLAUSE.AI AUDIT REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Meta Info
    tone = config.get("tone", "Standard")
    lang = config.get("language", "English")
    meta = doc.add_paragraph(f"Confidential Legal Analysis | Tone: {tone}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Helper to add sections
    def add_section(title, raw_text):
        if not raw_text: return
        
        # 1. Clean the text using the new robust function
        clean_text = clean_for_docx(raw_text)
        
        # 2. Add Heading
        h = doc.add_heading(title, level=1)
        run = h.runs[0]
        run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
        
        # 3. Add Body Text
        p = doc.add_paragraph(clean_text)
        p.paragraph_format.space_after = Pt(12)

    # --- CONTENT GENERATION ---

    # 1. Executive Synthesis
    if "synthesis" in results:
        add_section("EXECUTIVE SYNTHESIS", results["synthesis"].get("summary", ""))

    # 2. Agent Reports
    agents = [
        ("LEGAL RISK ANALYSIS", "legal"),
        ("FINANCIAL AUDIT", "finance"),
        ("COMPLIANCE CHECK", "compliance"),
        ("OPERATIONAL REVIEW", "operations")
    ]

    for title, key in agents:
        if key in results:
            add_section(title, results[key].get("summary", ""))

    # Save to memory
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer